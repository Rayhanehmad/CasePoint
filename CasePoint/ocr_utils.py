"""
OCR utilities for document text extraction using Tesseract
"""

import os
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
import pdfplumber
from typing import Tuple, Optional
import logging

logger = logging.getLogger(__name__)

class OCRService:
    """OCR service for extracting text from documents"""
    
    def __init__(self):
        """Initialize OCR service"""
        # Tesseract is installed via system package, no path config needed in Replit
        pass
    
    def extract_text_from_file(self, file_path: str, file_type: str) -> Tuple[Optional[str], Optional[float]]:
        """
        Extract text from file with confidence score
        
        Args:
            file_path: Path to the file
            file_type: Type of file (pdf, jpg, png, txt, etc.)
            
        Returns:
            Tuple of (extracted_text, confidence_score)
        """
        try:
            file_type = file_type.lower()
            
            if file_type == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_type in ['jpg', 'jpeg', 'png', 'tiff', 'tif', 'bmp']:
                return self._extract_from_image(file_path)
            elif file_type in ['txt', 'text']:
                return self._extract_from_txt(file_path)
            elif file_type in ['doc', 'docx']:
                # For DOCX, we'll use python-docx
                return self._extract_from_docx(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return None, None
                
        except Exception as e:
            logger.error(f"OCR extraction failed for {file_path}: {str(e)}")
            return None, None
    
    def _extract_from_pdf(self, file_path: str) -> Tuple[Optional[str], Optional[float]]:
        """Extract text from PDF file using pdfplumber, PyMuPDF, and OCR fallback"""
        
        try:
            # Method 1: Try pdfplumber first (best for born-digital PDFs with tables and layout)
            logger.info("Attempting text extraction with pdfplumber...")
            text_content = []
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_content.append(page_text)
                        logger.debug(f"pdfplumber extracted {len(page_text)} chars from page {page_num}")
            
            if text_content and len('\n'.join(text_content).strip()) > 100:
                combined_text = "\n\n".join(text_content)
                logger.info(f"pdfplumber successfully extracted {len(combined_text)} characters")
                return combined_text, 100.0  # pdfplumber extraction has 100% confidence
            
            # Method 2: Fallback to PyMuPDF if pdfplumber didn't work well
            logger.info("pdfplumber extraction insufficient, trying PyMuPDF...")
            text_content = []
            total_confidence = 0
            page_count = 0
            
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Try to extract text directly (for text-based PDFs)
                text = page.get_text()
                
                if text.strip():
                    text_content.append(text)
                    total_confidence += 100  # Direct text extraction has 100% confidence
                    page_count += 1
                else:
                    # Method 3: If no text found, use OCR on rendered page image
                    logger.info(f"No text found on page {page_num}, using OCR...")
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                    img_data = pix.tobytes("png")
                    
                    # Save temporary image for OCR
                    temp_img_path = f"/tmp/page_{page_num}_{os.getpid()}.png"
                    with open(temp_img_path, "wb") as f:
                        f.write(img_data)
                    
                    # Perform OCR
                    ocr_text, confidence = self._extract_from_image(temp_img_path)
                    if ocr_text:
                        text_content.append(ocr_text)
                        total_confidence += confidence or 0
                        page_count += 1
                    
                    # Clean up temp file
                    if os.path.exists(temp_img_path):
                        os.remove(temp_img_path)
            
            doc.close()
            
            if text_content:
                combined_text = "\n\n".join(text_content)
                avg_confidence = total_confidence / page_count if page_count > 0 else 0
                logger.info(f"PyMuPDF/OCR extracted {len(combined_text)} characters with {avg_confidence:.1f}% confidence")
                return combined_text, avg_confidence
            
            return None, None
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            return None, None
    
    def _extract_from_image(self, file_path: str) -> Tuple[Optional[str], Optional[float]]:
        """Extract text from image file using Tesseract OCR"""
        
        try:
            # Open and preprocess image
            image = Image.open(file_path)
            image = image.convert('RGB')
            
            # Configure OCR for English and Urdu (Pakistan languages)
            # PSM 6: Assume a single uniform block of text
            # OEM 3: Default, based on what is available
            config = '--oem 3 --psm 6'
            
            # Extract text with confidence data
            data = pytesseract.image_to_data(image, config=config, output_type=pytesseract.Output.DICT)
            
            # Filter out low confidence words and combine text
            text_parts = []
            confidences = []
            
            for i, conf in enumerate(data['conf']):
                try:
                    conf_val = int(conf)
                    if conf_val > 30:  # Only include words with confidence > 30%
                        word = data['text'][i].strip()
                        if word:
                            text_parts.append(word)
                            confidences.append(conf_val)
                except (ValueError, TypeError):
                    continue
            
            if text_parts:
                extracted_text = ' '.join(text_parts)
                avg_confidence = sum(confidences) / len(confidences) if confidences else 0
                logger.info(f"OCR extracted {len(text_parts)} words with avg confidence {avg_confidence:.1f}%")
                return extracted_text, avg_confidence
            
            return None, None
            
        except Exception as e:
            logger.error(f"Image OCR failed: {str(e)}")
            return None, None
    
    def _extract_from_txt(self, file_path: str) -> Tuple[Optional[str], Optional[float]]:
        """Extract text from plain text file"""
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return content, 100.0  # Plain text has 100% confidence
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                return content, 100.0
            except Exception as e:
                logger.error(f"Text file reading failed: {str(e)}")
                return None, None
        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            return None, None
    
    def _extract_from_docx(self, file_path: str) -> Tuple[Optional[str], Optional[float]]:
        """Extract text from DOCX file"""
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            if text_content:
                combined_text = '\n'.join(text_content)
                return combined_text, 100.0  # DOCX extraction has 100% confidence
            
            return None, None
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return None, None


# Create a global OCR service instance
ocr_service = OCRService()
