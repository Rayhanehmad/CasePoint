"""
DOCX Text Extraction Utility
Extracts comprehensive text from Microsoft Word documents including:
- Paragraphs
- Tables (all columns and rows)
- Headers and footers
- Multi-column sections
"""

from docx import Document
import logging

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path):
    """
    Extracts ALL text from a .docx file including:
    - Normal paragraphs
    - Table cells (all rows / columns)
    - Multi-column section text
    - Headers and footers
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Clean combined plain-text string with duplicates removed
        
    Raises:
        Exception: If document cannot be processed
    """
    try:
        doc = Document(file_path)
        output = []

        # 1. Extract normal paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                output.append(text)

        # 2. Extract text from tables (ALL columns + rows)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    # Join the row's text with a separator
                    output.append(" | ".join(row_text))

        # 3. Extract multi-column section text (Word "sections")
        for section in doc.sections:
            if hasattr(section, "header") and section.header:
                for paragraph in section.header.paragraphs:
                    header_text = paragraph.text.strip()
                    if header_text:
                        output.append(header_text)

            if hasattr(section, "footer") and section.footer:
                for paragraph in section.footer.paragraphs:
                    footer_text = paragraph.text.strip()
                    if footer_text:
                        output.append(footer_text)

        # Remove duplicates but keep order
        clean_output = []
        seen = set()
        for line in output:
            if line not in seen:
                clean_output.append(line)
                seen.add(line)

        return "\n".join(clean_output)
        
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {str(e)}")
        raise
