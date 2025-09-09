"""
Multi-Tenant SaaS Flask Application for KanoonPK Legal Research Platform
"""
import os
import uuid
import json
import datetime
import logging
from functools import wraps
from flask import Flask, request, jsonify, render_template, redirect, url_for, flash, g, session
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from flask_jwt_extended import JWTManager, jwt_required, create_access_token, get_jwt_identity, get_jwt
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.utils import secure_filename
from werkzeug.security import check_password_hash
from openai import OpenAI
import chromadb
from chromadb.utils import embedding_functions
from PyPDF2 import PdfReader
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors

# Import our models
from models import (
    db, Tenant, User, Subscription, UsageMetric, LegalDocument, QueryHistory, LegalWorkspace,
    create_tenant_schema, switch_tenant_schema, get_tenant_from_subdomain, record_usage, PLAN_LIMITS
)

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# =============================================================================
# FLASK APP INITIALIZATION
# =============================================================================

app = Flask(__name__)

# Configuration
app.config['SECRET_KEY'] = os.environ.get('SESSION_SECRET', 'dev-secret-key')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_recycle': 300,
    'pool_pre_ping': True,
}

# File upload settings
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# JWT Configuration
app.config['JWT_SECRET_KEY'] = os.environ.get('JWT_SECRET_KEY', 'jwt-secret-string')
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = datetime.timedelta(hours=24)

# Initialize extensions
db.init_app(app)
migrate = Migrate(app, db)
jwt = JWTManager(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'auth.login'

# Initialize OpenAI
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Initialize ChromaDB with tenant isolation
chroma_client = chromadb.PersistentClient(path="chroma_db")
embedding_func = embedding_functions.OpenAIEmbeddingFunction(
    api_key=os.getenv("OPENAI_API_KEY"),
    model_name="text-embedding-3-small"
)

# =============================================================================
# USER LOADER AND JWT CALLBACKS
# =============================================================================

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

@jwt.user_identity_loader
def user_identity_lookup(user):
    return user.id

@jwt.user_lookup_loader
def user_lookup_callback(_jwt_header, jwt_data):
    identity = jwt_data["sub"]
    return User.query.filter_by(id=identity).one_or_none()

# =============================================================================
# TENANT MIDDLEWARE
# =============================================================================

@app.before_request
def load_tenant_context():
    """Load tenant context from subdomain and set up database schema"""
    # Skip tenant loading for static files and certain endpoints
    if request.endpoint and (request.endpoint.startswith('static') or 
                            request.endpoint in ['auth.register_tenant', 'health']):
        return
    
    # Extract subdomain from host
    host = request.host.lower()
    subdomain = None
    
    # Handle localhost and development
    if 'localhost' in host or '127.0.0.1' in host:
        # For development, use subdomain from query param or session
        subdomain = request.args.get('tenant') or session.get('tenant_subdomain')
    else:
        # Extract subdomain from host
        parts = host.split('.')
        if len(parts) > 2 and parts[0] not in ['www', 'api', 'admin']:
            subdomain = parts[0]
    
    if subdomain:
        # Load tenant from database
        tenant = get_tenant_from_subdomain(subdomain)
        if tenant:
            g.tenant = tenant
            # Switch to tenant's schema
            switch_tenant_schema(tenant.id)
            # Store in session for development
            session['tenant_subdomain'] = subdomain
        else:
            # Tenant not found
            if request.endpoint not in ['auth.login', 'auth.register']:
                return jsonify({'error': 'Tenant not found'}), 404

def require_tenant(f):
    """Decorator to ensure request has valid tenant context"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not hasattr(g, 'tenant'):
            return jsonify({'error': 'No tenant context'}), 400
        return f(*args, **kwargs)
    return decorated_function

def require_plan_feature(feature_name):
    """Decorator to check if tenant's plan includes specific feature"""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if not hasattr(g, 'tenant'):
                return jsonify({'error': 'No tenant context'}), 400
            
            plan_features = PLAN_LIMITS.get(g.tenant.plan, {}).get('features', [])
            if feature_name not in plan_features and 'all_features' not in plan_features:
                return jsonify({'error': f'Feature "{feature_name}" not available in your plan'}), 403
            
            return f(*args, **kwargs)
        return decorated_function
    return decorator

def require_usage_limit(action_type):
    """Decorator to check usage limits before allowing action"""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if not hasattr(g, 'tenant'):
                return jsonify({'error': 'No tenant context'}), 400
            
            if not g.tenant.can_perform_action(action_type):
                return jsonify({'error': f'Usage limit exceeded for {action_type}'}), 429
            
            return f(*args, **kwargs)
        return decorated_function
    return decorator

# =============================================================================
# LEGAL RESEARCH ENGINE (Enhanced)
# =============================================================================

class EnhancedLegalSearchEngine:
    """Enhanced legal search engine with Pakistan law specific features"""
    
    PAKISTAN_JURISDICTIONS = [
        'Supreme Court of Pakistan',
        'Federal Shariat Court',
        'Lahore High Court',
        'Karachi High Court (Sindh)',
        'Peshawar High Court',
        'Quetta High Court (Balochistan)',
        'Islamabad High Court'
    ]
    
    LEGAL_AREAS = [
        'Constitutional Law', 'Criminal Law', 'Civil Law', 'Commercial Law',
        'Islamic Law', 'Administrative Law', 'Labor Law', 'Family Law',
        'Property Law', 'Contract Law', 'Corporate Law', 'Tax Law'
    ]
    
    def __init__(self, tenant_id):
        self.tenant_id = tenant_id
        self.collection_name = f"tenant_{tenant_id}_legal_docs"
        self.collection = self._get_or_create_collection()
    
    def _get_or_create_collection(self):
        """Get or create ChromaDB collection for tenant"""
        try:
            collection = chroma_client.get_collection(self.collection_name)
        except:
            collection = chroma_client.create_collection(
                self.collection_name, 
                embedding_function=embedding_func
            )
        return collection
    
    def advanced_search(self, query, filters=None, n_results=10):
        """Enhanced search with Pakistan law specific filters"""
        filters = filters or {}
        
        # Build ChromaDB where conditions
        where_conditions = {}
        
        if filters.get('jurisdiction'):
            where_conditions['jurisdiction'] = {'$in': filters['jurisdiction']}
        if filters.get('legal_area'):
            where_conditions['legal_area'] = {'$in': filters['legal_area']}
        if filters.get('document_type'):
            where_conditions['document_type'] = {'$in': filters['document_type']}
        if filters.get('court_level'):
            where_conditions['court_level'] = {'$in': filters['court_level']}
        if filters.get('date_range'):
            # Add date filtering if metadata contains dates
            pass
        
        try:
            # Perform vector search
            results = self.collection.query(
                query_texts=[query],
                n_results=n_results,
                where=where_conditions if where_conditions else None,
                include=['documents', 'metadatas', 'distances']
            )
            
            return self._process_search_results(results, query)
            
        except Exception as e:
            logger.error(f"Search error: {e}")
            return {
                'documents': [],
                'citations': [],
                'metadata': [],
                'error': str(e)
            }
    
    def _process_search_results(self, results, original_query):
        """Process and enhance search results"""
        if not results['documents'] or not results['documents'][0]:
            return {
                'documents': [],
                'citations': [],
                'metadata': [],
                'message': 'No relevant documents found'
            }
        
        processed_results = {
            'documents': [],
            'citations': [],
            'metadata': [],
            'confidence_scores': []
        }
        
        for i, doc in enumerate(results['documents'][0]):
            metadata = results['metadatas'][0][i] if i < len(results['metadatas'][0]) else {}
            distance = results['distances'][0][i] if i < len(results['distances'][0]) else 1.0
            confidence = max(0, 1 - distance)  # Convert distance to confidence score
            
            processed_results['documents'].append(doc)
            processed_results['metadata'].append(metadata)
            processed_results['confidence_scores'].append(confidence)
            
            # Extract citation if available
            citation = metadata.get('citation', metadata.get('source', ''))
            if citation and citation not in processed_results['citations']:
                processed_results['citations'].append(citation)
        
        return processed_results
    
    def extract_legal_citations(self, text):
        """Extract Pakistan legal citations from text"""
        import re
        
        # Pakistan legal citation patterns
        patterns = {
            'pld': r'(\d{4})\s+(PLD)\s+(\d+)\s+(\w+)',
            'scmr': r'(\d{4})\s+(SCMR)\s+(\d+)',
            'clr': r'(\d{4})\s+(CLR)\s+(\d+)',
            'mld': r'(\d{4})\s+(MLD)\s+(\d+)\s+(\w+)',
            'ylr': r'(\d{4})\s+(YLR)\s+(\d+)',
            'plc': r'(\d{4})\s+(PLC)\s+(\d+)'
        }
        
        citations = []
        for citation_type, pattern in patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                citations.append({
                    'type': citation_type.upper(),
                    'year': match[0],
                    'reporter': match[1],
                    'volume': match[2],
                    'court': match[3] if len(match) > 3 else '',
                    'full_citation': ' '.join(match)
                })
        
        return citations
    
    def analyze_precedents(self, case_text, n_results=5):
        """Find similar cases and legal precedents"""
        try:
            # Search for similar cases
            results = self.collection.query(
                query_texts=[case_text],
                n_results=n_results,
                where={'document_type': 'case_law'},
                include=['documents', 'metadatas', 'distances']
            )
            
            precedents = []
            if results['documents'] and results['documents'][0]:
                for i, doc in enumerate(results['documents'][0]):
                    metadata = results['metadatas'][0][i]
                    distance = results['distances'][0][i]
                    similarity = max(0, 1 - distance)
                    
                    precedents.append({
                        'case_name': metadata.get('case_name', ''),
                        'citation': metadata.get('citation', ''),
                        'court': metadata.get('jurisdiction', ''),
                        'similarity_score': round(similarity, 3),
                        'relevant_text': doc[:300] + "..." if len(doc) > 300 else doc,
                        'legal_area': metadata.get('legal_area', [])
                    })
            
            return precedents
            
        except Exception as e:
            logger.error(f"Precedent analysis error: {e}")
            return []

# =============================================================================
# DOCUMENT MANAGEMENT
# =============================================================================

class TenantDocumentManager:
    """Enhanced document management for tenants"""
    
    def __init__(self, tenant_id):
        self.tenant_id = tenant_id
        self.tenant = Tenant.query.get(tenant_id)
        self.search_engine = EnhancedLegalSearchEngine(tenant_id)
        
    def upload_document(self, file, user_id, metadata=None):
        """Upload and process legal document"""
        metadata = metadata or {}
        
        try:
            # Check storage limits
            if not self.tenant.can_perform_action('document_upload'):
                raise Exception("Document upload limit exceeded")
            
            # Secure filename
            filename = secure_filename(file.filename)
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            unique_filename = f"{timestamp}_{filename}"
            
            # Save file
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
            file.save(file_path)
            
            # Extract text content
            text_content = self._extract_text_from_file(file_path, file.filename)
            
            # Analyze and categorize document
            doc_analysis = self._analyze_legal_document(text_content)
            
            # Create database record
            doc_record = LegalDocument(
                filename=unique_filename,
                original_filename=file.filename,
                file_size=os.path.getsize(file_path),
                file_path=file_path,
                document_type=doc_analysis.get('document_type', 'general'),
                legal_area=doc_analysis.get('legal_areas', []),
                jurisdiction=doc_analysis.get('jurisdiction', ''),
                court_level=doc_analysis.get('court_level', ''),
                extracted_citations=doc_analysis.get('citations', []),
                extracted_entities=doc_analysis.get('entities', []),
                upload_user_id=user_id,
                processing_status='processed'
            )
            
            # Add to ChromaDB
            chunks_added = self._add_to_vector_store(text_content, doc_record, doc_analysis)
            doc_record.total_chunks = chunks_added
            
            # Save to database
            db.session.add(doc_record)
            db.session.commit()
            
            # Record usage
            record_usage(self.tenant_id, user_id, 'document_upload', 1, {
                'filename': file.filename,
                'file_size': doc_record.file_size,
                'document_type': doc_record.document_type
            })
            
            return {
                'success': True,
                'document_id': doc_record.id,
                'chunks_processed': chunks_added,
                'document_type': doc_record.document_type,
                'legal_areas': doc_record.legal_area
            }
            
        except Exception as e:
            logger.error(f"Document upload error: {e}")
            # Clean up file if it was saved
            if 'file_path' in locals() and os.path.exists(file_path):
                os.remove(file_path)
            return {
                'success': False,
                'error': str(e)
            }
    
    def _extract_text_from_file(self, file_path, filename):
        """Extract text from uploaded file"""
        try:
            if filename.lower().endswith('.pdf'):
                return self._extract_from_pdf(file_path)
            elif filename.lower().endswith('.docx'):
                return self._extract_from_docx(file_path)
            elif filename.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            else:
                raise Exception("Unsupported file format")
        except Exception as e:
            raise Exception(f"Text extraction failed: {e}")
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF"""
        try:
            reader = PdfReader(file_path)
            text_content = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
            return "\n".join(text_content)
        except Exception as e:
            raise Exception(f"PDF extraction error: {e}")
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        except Exception as e:
            raise Exception(f"DOCX extraction error: {e}")
    
    def _analyze_legal_document(self, text_content):
        """Analyze document content for legal metadata"""
        analysis = {
            'document_type': 'general',
            'legal_areas': [],
            'jurisdiction': '',
            'court_level': '',
            'citations': [],
            'entities': []
        }
        
        text_lower = text_content.lower()
        
        # Document type classification
        type_keywords = {
            'case_law': ['judgment', 'order', 'decision', 'ruling', 'appeal', 'writ'],
            'statute': ['act', 'ordinance', 'regulation', 'code', 'law'],
            'contract': ['agreement', 'contract', 'deed', 'instrument'],
            'pleading': ['petition', 'application', 'appeal', 'writ petition'],
            'opinion': ['opinion', 'advice', 'memorandum', 'brief']
        }
        
        type_scores = {}
        for doc_type, keywords in type_keywords.items():
            score = sum(text_lower.count(keyword) for keyword in keywords)
            type_scores[doc_type] = score
        
        if type_scores:
            analysis['document_type'] = max(type_scores, key=type_scores.get)
        
        # Extract legal citations
        analysis['citations'] = self.search_engine.extract_legal_citations(text_content)
        
        # Identify jurisdiction
        for jurisdiction in EnhancedLegalSearchEngine.PAKISTAN_JURISDICTIONS:
            if jurisdiction.lower() in text_lower:
                analysis['jurisdiction'] = jurisdiction
                break
        
        # Identify legal areas
        for area in EnhancedLegalSearchEngine.LEGAL_AREAS:
            if area.lower() in text_lower:
                analysis['legal_areas'].append(area)
        
        return analysis
    
    def _add_to_vector_store(self, text_content, doc_record, analysis):
        """Add document chunks to ChromaDB"""
        # Create chunks
        chunk_size = 1000
        overlap = 200
        chunks = []
        
        for i in range(0, len(text_content), chunk_size - overlap):
            chunk = text_content[i:i + chunk_size]
            if chunk.strip():
                chunk_metadata = {
                    'source': doc_record.original_filename,
                    'document_id': str(uuid.uuid4()),  # Will be updated with actual ID
                    'chunk_index': len(chunks),
                    'document_type': analysis.get('document_type', 'general'),
                    'legal_area': analysis.get('legal_areas', []),
                    'jurisdiction': analysis.get('jurisdiction', ''),
                    'court_level': analysis.get('court_level', ''),
                    'citation': doc_record.original_filename,
                    'upload_date': datetime.datetime.now().isoformat()
                }
                chunks.append({
                    'text': chunk,
                    'metadata': chunk_metadata
                })
        
        if not chunks:
            raise Exception("No valid text chunks created")
        
        # Add to ChromaDB
        try:
            self.search_engine.collection.add(
                documents=[chunk['text'] for chunk in chunks],
                metadatas=[chunk['metadata'] for chunk in chunks],
                ids=[f"{doc_record.filename}_chunk_{i}_{str(uuid.uuid4())[:8]}" for i in range(len(chunks))]
            )
            return len(chunks)
        except Exception as e:
            raise Exception(f"Vector store error: {e}")

# Continue with routes in next part...