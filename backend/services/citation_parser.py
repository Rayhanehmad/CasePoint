"""
Pakistani Legal Citation Parser and Extractor
Recognizes and extracts citations in Pakistani legal formats
"""

import re
import logging
from typing import List, Dict, Tuple, Optional
from docx import Document

logger = logging.getLogger(__name__)

class CitationParser:
    """Parser for Pakistani legal citation formats"""
    
    # Pakistani legal citation patterns
    CITATION_PATTERNS = [
        # PLD format: PLD 2024 SC 1, PLD 2024 Karachi 1
        r'PLD\s+(\d{4})\s+([A-Za-z\s]+?)\s+(\d+)',
        
        # Standard format: YEAR CODE NUMBER (e.g., 2025 CLC 1, 2025 SCMR 1)
        r'(\d{4})\s+(CLC|CLD|MLD|PCrLJ|PLC|PTD|SCMR|YLR)\s+(\d+)',
        
        # CLC format variations
        r'CLC\s+(\d{4})\s+(\d+)',
        
        # General format with court: CODE YEAR COURT NUMBER
        r'(PLD|CLC|CLD|MLD|PCrLJ|PLC|PTD|SCMR|YLR)\s+(\d{4})\s+([A-Za-z\s]+?)\s+(\d+)',
    ]
    
    # Compile all patterns
    COMPILED_PATTERNS = [re.compile(pattern, re.IGNORECASE) for pattern in CITATION_PATTERNS]
    
    # Citation codes
    VALID_CODES = ['PLD', 'CLC', 'CLD', 'MLD', 'PCrLJ', 'PLC', 'PTD', 'SCMR', 'YLR']
    
    # Court abbreviations
    COURT_MAPPINGS = {
        'SC': 'Supreme Court of Pakistan',
        'FSC': 'Federal Shariat Court',
        'Karachi': 'Karachi High Court',
        'Lahore': 'Lahore High Court',
        'Islamabad': 'Islamabad High Court',
        'Peshawar': 'Peshawar High Court',
        'Quetta': 'Quetta High Court',
    }
    
    def __init__(self):
        """Initialize citation parser"""
        pass
    
    def find_all_citations(self, text: str) -> List[Dict[str, any]]:
        """
        Find all Pakistani legal citations in text
        
        Returns:
            List of dicts with: {
                'citation': 'PLD 2024 SC 1',
                'code': 'PLD',
                'year': 2024,
                'court': 'Supreme Court of Pakistan',
                'number': 1,
                'position': (start, end)
            }
        """
        citations = []
        seen_citations = set()  # Avoid duplicates
        
        for pattern in self.COMPILED_PATTERNS:
            for match in pattern.finditer(text):
                citation_info = self._extract_citation_info(match, text)
                if citation_info and citation_info['citation'] not in seen_citations:
                    citations.append(citation_info)
                    seen_citations.add(citation_info['citation'])
        
        # Sort by position in text
        citations.sort(key=lambda x: x['position'][0])
        return citations
    
    def _extract_citation_info(self, match, text: str) -> Optional[Dict]:
        """Extract structured citation information from regex match"""
        try:
            groups = match.groups()
            matched_text = match.group(0)
            
            # Determine citation format
            if matched_text.upper().startswith('PLD'):
                # PLD format
                if len(groups) == 3:
                    year = int(groups[0])
                    court_raw = groups[1].strip()
                    number = int(groups[2])
                    code = 'PLD'
                else:
                    return None
            elif groups[0].isdigit():
                # Year-first format (e.g., 2025 CLC 1)
                year = int(groups[0])
                code = groups[1].upper()
                number = int(groups[2])
                court_raw = None
            elif groups[0].upper() in self.VALID_CODES:
                # Code-first format
                code = groups[0].upper()
                if len(groups) == 2:
                    year = int(groups[1])
                    number = int(groups[2]) if len(groups) > 2 else 1
                    court_raw = None
                elif len(groups) == 4:
                    year = int(groups[1])
                    court_raw = groups[2].strip()
                    number = int(groups[3])
                else:
                    return None
            else:
                return None
            
            # Map court name
            court = self._map_court_name(court_raw) if court_raw else self._default_court_for_code(code)
            
            return {
                'citation': matched_text,
                'code': code,
                'year': year,
                'court': court,
                'number': number,
                'position': match.span()
            }
            
        except (ValueError, IndexError) as e:
            logger.debug(f"Could not parse citation: {match.group(0)} - {str(e)}")
            return None
    
    def _map_court_name(self, court_raw: str) -> str:
        """Map court abbreviation to full name"""
        court_raw = court_raw.strip()
        return self.COURT_MAPPINGS.get(court_raw, court_raw + ' Court')
    
    def _default_court_for_code(self, code: str) -> str:
        """Get default court for citation code"""
        defaults = {
            'PLD': 'Supreme Court of Pakistan',
            'SCMR': 'Supreme Court of Pakistan',
            'CLC': 'High Court',
            'CLD': 'High Court',
            'MLD': 'High Court',
            'PCrLJ': 'High Court',
            'PLC': 'Labour Court',
            'PTD': 'Tax Court',
            'YLR': 'High Court',
        }
        return defaults.get(code, 'Court of Pakistan')
    
    def split_document_by_citations(self, text: str) -> List[Dict[str, any]]:
        """
        Split document into individual citation blocks using improved lookahead regex
        that keeps internal references intact
        
        Returns:
            List of dicts with: {
                'citation': 'PLD 2024 SC 1',
                'code': 'PLD',
                'year': 2024,
                'court': 'Supreme Court',
                'number': 1,
                'text': 'Full text content for this citation...'
            }
        """
        return self.split_document_by_citations_improved(text)
    
    def split_document_by_citations_improved(self, text: str) -> List[Dict[str, any]]:
        """
        Improved citation extraction that identifies main citations and captures their full blocks
        Uses paragraph boundaries to avoid breaking at internal references
        """
        # Split into paragraphs first (double newlines typically separate cases)
        paragraphs = re.split(r'\n\s*\n', text)
        
        citation_blocks_raw = []
        current_block = ""
        current_citation = None
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Check if this paragraph starts with a main citation
            # Main citations typically appear at the start of a line
            first_line = para.split('\n')[0].strip()
            
            # Try to match citation at the beginning of the paragraph
            main_citation = None
            for pattern in self.COMPILED_PATTERNS:
                match = pattern.match(first_line)
                if match:
                    citation_info = self._extract_citation_info(match, first_line)
                    if citation_info:
                        main_citation = citation_info
                        break
            
            if main_citation:
                # Save previous block if exists
                if current_block and current_citation:
                    citation_blocks_raw.append({
                        'text': current_block,
                        'citation_info': current_citation
                    })
                
                # Start new block
                current_block = para
                current_citation = main_citation
            else:
                # Add to current block
                if current_block:
                    current_block += "\n\n" + para
                else:
                    # Orphan paragraph - try to find any citation in it
                    for pattern in self.COMPILED_PATTERNS:
                        match = pattern.search(para)
                        if match:
                            citation_info = self._extract_citation_info(match, para)
                            if citation_info:
                                current_block = para
                                current_citation = citation_info
                                break
        
        # Add last block
        if current_block and current_citation:
            citation_blocks_raw.append({
                'text': current_block,
                'citation_info': current_citation
            })
        
        logger.info(f"Found {len(citation_blocks_raw)} citation blocks")
        
        # Convert to final format
        citation_blocks = []
        for i, block_data in enumerate(citation_blocks_raw, 1):
            citation_info = block_data['citation_info']
            block_text = block_data['text']
            
            citation_blocks.append({
                'citation': citation_info['citation'],
                'code': citation_info['code'],
                'year': citation_info['year'],
                'court': citation_info['court'],
                'number': citation_info.get('number', 0),
                'text': block_text,
                'text_length': len(block_text)
            })
            
            # Debug output with clear markers
            logger.debug(f"\n===== START CITATION {i} =====")
            logger.debug(f"[CITATION] {citation_info['citation']}")
            logger.debug(f"[PREVIEW] {block_text[:100]}...")
            logger.debug(f"[LENGTH] {len(block_text)} chars")
            logger.debug(f"===== END CITATION {i} =====\n")
        
        logger.info(f"Extracted {len(citation_blocks)} complete citation blocks")
        return citation_blocks
    
    def _merge_small_fragments(self, blocks: List[str]) -> List[str]:
        """
        Merge small fragments (< 200 chars) into larger blocks to fix false breaks
        """
        cleaned = []
        buffer = ""
        
        for block in blocks:
            block_stripped = block.strip()
            
            # If block is too small, add to buffer
            if len(block_stripped) < 200:
                buffer += " " + block_stripped
            else:
                # This is a substantial block
                if buffer:
                    # Merge buffer with this block
                    cleaned.append(buffer + " " + block_stripped)
                    buffer = ""
                else:
                    # Just add the block
                    cleaned.append(block_stripped)
        
        # Add any remaining buffer
        if buffer.strip():
            cleaned.append(buffer.strip())
        
        return cleaned
    
    def _extract_citation_from_block(self, block_text: str) -> Optional[Dict]:
        """
        Extract the main citation identifier from a block of text
        """
        # Look for the first citation pattern in the block
        for pattern in self.COMPILED_PATTERNS:
            match = pattern.search(block_text)
            if match:
                citation_info = self._extract_citation_info(match, block_text)
                if citation_info:
                    return citation_info
        
        return None
    
    def extract_title_from_text(self, citation_text: str, citation_code: str) -> Optional[str]:
        """
        Extract case title from citation text
        Usually appears right after the citation line
        """
        lines = citation_text.split('\n')
        
        # Skip the citation line itself
        for i, line in enumerate(lines):
            line = line.strip()
            
            # Skip empty lines and citation lines
            if not line or citation_code in line:
                continue
            
            # Look for common title patterns
            # Typically: "PARTY v. PARTY" or "In the matter of..."
            if ' v. ' in line or ' vs. ' in line or ' V. ' in line:
                return line[:200]  # Limit title length
            
            # If line has reasonable length and doesn't look like a citation
            if 20 < len(line) < 200 and not any(code in line for code in self.VALID_CODES):
                return line
        
        return None
    
    def split_document_by_headings(self, filepath: str) -> List[Dict[str, any]]:
        """
        Extract citations from DOCX using Heading 1 styles
        
        This method is faster and more accurate for properly formatted DOCX files
        where each citation is marked with Heading 1 style
        
        Args:
            filepath: Path to the DOCX file
            
        Returns:
            List of dicts with citation information and text content
        """
        try:
            doc = Document(filepath)
            citations = []
            current_title = None
            current_text = ""
            
            for para in doc.paragraphs:
                style = para.style.name
                
                # Check if this is a heading (citation identifier)
                if style.startswith("Heading 1"):
                    # Save previous citation if exists
                    if current_title and current_text.strip():
                        citation_data = self._parse_heading_citation(current_title, current_text.strip())
                        if citation_data:
                            citations.append(citation_data)
                    
                    # Start new citation
                    current_title = para.text.strip()
                    current_text = ""
                else:
                    # Add to current citation body
                    if para.text.strip():
                        current_text += para.text + " "
            
            # Append last citation
            if current_title and current_text.strip():
                citation_data = self._parse_heading_citation(current_title, current_text.strip())
                if citation_data:
                    citations.append(citation_data)
            
            logger.info(f"Extracted {len(citations)} citations using heading-based method")
            return citations
            
        except Exception as e:
            logger.error(f"Error extracting citations from headings: {str(e)}")
            return []
    
    def _parse_heading_citation(self, heading: str, body_text: str) -> Optional[Dict]:
        """
        Parse citation information from heading text
        
        Args:
            heading: The Heading 1 text (e.g., "2003 MLD 1077")
            body_text: The full text content under this heading
            
        Returns:
            Dict with citation info and text, or None if parsing fails
        """
        # Try to extract citation from heading using patterns
        citation_info = None
        for pattern in self.COMPILED_PATTERNS:
            match = pattern.search(heading)
            if match:
                citation_info = self._extract_citation_info(match, heading)
                if citation_info:
                    break
        
        if not citation_info:
            # If no pattern match, try to find citation in the body text
            for pattern in self.COMPILED_PATTERNS:
                match = pattern.search(body_text[:200])  # Check first 200 chars
                if match:
                    citation_info = self._extract_citation_info(match, body_text)
                    if citation_info:
                        break
        
        if citation_info:
            return {
                'citation': citation_info['citation'],
                'code': citation_info['code'],
                'year': citation_info['year'],
                'court': citation_info['court'],
                'number': citation_info.get('number', 0),
                'text': body_text,
                'text_length': len(body_text)
            }
        
        # Fallback: use heading as citation if no pattern found
        logger.warning(f"Could not parse citation pattern from heading: {heading}")
        return {
            'citation': heading[:100],  # Use heading as citation
            'code': 'UNKNOWN',
            'year': None,
            'court': 'Unknown Court',
            'number': 0,
            'text': body_text,
            'text_length': len(body_text)
        }
    
    def auto_detect_and_split(self, filepath: str, text: str = None, method: str = 'auto') -> Tuple[List[Dict], str]:
        """
        Auto-detect the best extraction method and split document
        
        Args:
            filepath: Path to the DOCX file
            text: Plain text content (optional, used for pattern-based method)
            method: 'auto', 'heading', or 'pattern'
            
        Returns:
            Tuple of (citations list, method_used)
        """
        if method == 'heading':
            citations = self.split_document_by_headings(filepath)
            return citations, 'heading-based'
        
        if method == 'pattern':
            if not text:
                # Read text from file
                from docx import Document
                doc = Document(filepath)
                text = '\n'.join([para.text for para in doc.paragraphs])
            citations = self.split_document_by_citations(text)
            return citations, 'pattern-based'
        
        # Auto-detect method
        # Try heading-based first
        citations_heading = self.split_document_by_headings(filepath)
        
        # Check if heading-based found meaningful citations
        if len(citations_heading) >= 3:  # At least 3 citations with headings
            logger.info(f"Auto-detected heading-based method ({len(citations_heading)} citations)")
            return citations_heading, 'heading-based (auto-detected)'
        
        # Fall back to pattern-based
        if not text:
            doc = Document(filepath)
            text = '\n'.join([para.text for para in doc.paragraphs])
        
        citations_pattern = self.split_document_by_citations(text)
        logger.info(f"Auto-detected pattern-based method ({len(citations_pattern)} citations)")
        return citations_pattern, 'pattern-based (auto-detected)'


# Global instance
citation_parser = CitationParser()
