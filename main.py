import os
import uuid
import json
import datetime
import csv
from flask import Flask, request, jsonify, render_template_string, send_file
from werkzeug.utils import secure_filename
from openai import OpenAI
import chromadb
from chromadb.utils import embedding_functions
from PyPDF2 import PdfReader
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors

# ----------------------------
# Flask setup
# ----------------------------
app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = "uploads"
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

# ----------------------------
# OpenAI client
# ----------------------------
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ----------------------------
# ChromaDB setup
# ----------------------------
chroma_client = chromadb.PersistentClient(path="chroma_db")
embedding_func = embedding_functions.OpenAIEmbeddingFunction(
    api_key=os.getenv("OPENAI_API_KEY"),
    model_name="text-embedding-3-small"
)
try:
    collection = chroma_client.get_collection("kanoonpk")
except:
    collection = chroma_client.create_collection("kanoonpk", embedding_function=embedding_func)

LEGAL_SYSTEM_PROMPT = """
You are KanoonPK, an AI Legal Research Assistant specialized in Pakistan law.
Answer based on the provided documents and Pakistan's laws, case references.
Always provide citations if available from the documents.
"""

GENERAL_LEGAL_PROMPT = """
You are KanoonPK, an AI Legal Research Assistant specialized in Pakistan law.
Since no specific documents were found for this query, use your knowledge of Pakistan law to provide accurate information.
Focus on:
- Constitution of Pakistan 1973
- Pakistan Penal Code
- Civil Procedure Code
- Criminal Procedure Code
- Contract Act 1872
- Companies Act 2017
- Other relevant Pakistan legal statutes

Always mention that this answer is based on general legal knowledge and suggest uploading specific documents for more detailed analysis.
"""

# ----------------------------
# File processing
# ----------------------------
def get_uploaded_documents():
    """Get list of uploaded documents from ChromaDB"""
    try:
        # Get all documents from collection
        all_docs = collection.get()
        
        # Group by source file to avoid duplicates from chunks
        documents = {}
        for i, metadata in enumerate(all_docs['metadatas']):
            source = metadata.get('source', 'Unknown')
            if source not in documents:
                documents[source] = {
                    'filename': source,
                    'citation': metadata.get('citation', ''),
                    'year': metadata.get('year', ''),
                    'page': metadata.get('page', ''),
                    'court': metadata.get('court', ''),
                    'chunks': 1
                }
            else:
                documents[source]['chunks'] += 1
        
        return list(documents.values())
    except Exception as e:
        print(f"Error getting documents: {e}")
        return []

def delete_document(filename):
    """Delete document from both filesystem and ChromaDB"""
    try:
        # Delete from ChromaDB
        all_docs = collection.get()
        ids_to_delete = []
        
        for i, metadata in enumerate(all_docs['metadatas']):
            if metadata.get('source') == filename:
                ids_to_delete.append(all_docs['ids'][i])
        
        if ids_to_delete:
            collection.delete(ids=ids_to_delete)
        
        # Delete from filesystem
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        if os.path.exists(file_path):
            os.remove(file_path)
        
        return True
    except Exception as e:
        print(f"Error deleting document {filename}: {e}")
        return False
def extract_text_from_pdf(path):
    reader = PdfReader(path)
    return "\n".join([page.extract_text() or "" for page in reader.pages])

def extract_text_from_docx(path):
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

def save_to_collection(file_path, filename, citation="", year="", page="", court=""):
    if filename.lower().endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
    elif filename.lower().endswith(".docx"):
        text = extract_text_from_docx(file_path)
    elif filename.lower().endswith((".txt", ".text")):
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    elif filename.lower().endswith((".jpg", ".jpeg", ".png")):
        # For image files, we'll store the filename as text
        text = f"Image file: {filename}"
    else:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

    chunks = [text[i:i+1000] for i in range(0, len(text), 1000)]
    ids = [str(uuid.uuid4()) for _ in chunks]
    metas = [{
        "source": filename, 
        "citation": citation or filename,
        "year": year,
        "page": page,
        "court": court
    }] * len(chunks)
    collection.add(documents=chunks, metadatas=metas, ids=ids)

# ----------------------------
# Routes
# ----------------------------

@app.route("/")
def home():
    return render_template_string("""
    <!DOCTYPE html>
    <html>
    <head>
      <title>KanoonPK - AI Legal Research</title>
      <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%); min-height: 100vh; font-size: 14px; }
        .chat-container { max-width: 100%; margin: 0; padding: 10px; }
        .header { background: white; padding: 8px 12px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); margin-bottom: 10px; display: flex; align-items: center; justify-content: space-between; }
        .header-left { display: flex; align-items: center; }
        .logo { width: 50px; height: auto; margin-right: 10px; }
        .header-text h1 { color: #2bc77a; font-size: 16px; font-weight: 700; margin: 0; }
        .header-text .tagline { color: #666; font-size: 11px; margin: 0; }
        .admin-link { color: #4dd0b7; text-decoration: none; font-size: 11px; font-weight: 600; padding: 5px 10px; border: 1px solid #4dd0b7; border-radius: 15px; transition: all 0.3s; }
        .admin-link:hover { background: #4dd0b7; color: white; }
        .user-profile { display: flex; align-items: center; gap: 8px; }
        .user-avatar { width: 35px; height: 35px; border-radius: 50%; border: 2px solid #4dd0b7; }
        .user-info { display: flex; flex-direction: column; }
        .user-name { font-size: 12px; font-weight: 600; color: #2bc77a; margin: 0; }
        .user-status { font-size: 10px; color: #666; margin: 0; }
        .chat-actions { display: flex; gap: 8px; margin-bottom: 10px; }
        .chat-action-btn { padding: 6px 12px; border: none; border-radius: 15px; cursor: pointer; font-size: 11px; font-weight: 600; transition: all 0.3s; }
        .delete-chat-btn { background: #ff6b6b; color: white; }
        .delete-chat-btn:hover { background: #ee5a24; }
        .select-mode-btn { background: #6c5ce7; color: white; }
        .select-mode-btn:hover { background: #5a4fcf; }
        .select-mode-btn.active { background: #5a4fcf; }
        .bulk-actions { display: none; gap: 5px; padding: 8px; background: white; border-radius: 10px; margin-bottom: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
        .bulk-actions.active { display: flex; }
        .bulk-action-btn { padding: 6px 12px; border: none; border-radius: 10px; cursor: pointer; font-size: 11px; font-weight: 600; }
        .delete-selected-btn { background: #ff6b6b; color: white; }
        .forward-selected-btn { background: #4dd0b7; color: white; }
        .cancel-selection-btn { background: #6c757d; color: white; }
        .export-actions { display: flex; gap: 8px; margin-bottom: 10px; }
        .export-btn { padding: 6px 12px; border: none; border-radius: 15px; cursor: pointer; font-size: 11px; font-weight: 600; transition: all 0.3s; }
        .export-txt-btn { background: #28a745; color: white; }
        .export-pdf-btn { background: #ff6b6b; color: white; }
        .export-json-btn { background: #6f42c1; color: white; }
        .export-dropdown { position: relative; display: inline-block; }
        .export-menu { display: none; position: absolute; background: white; min-width: 160px; box-shadow: 0 4px 20px rgba(0,0,0,0.1); border-radius: 10px; z-index: 1000; top: 100%; left: 0; }
        .export-menu.show { display: block; }
        .export-option { padding: 10px 15px; cursor: pointer; font-size: 12px; border-bottom: 1px solid #f0f0f0; }
        .export-option:hover { background: #f8f9fa; }
        .export-option:last-child { border-bottom: none; }
        .msg { padding: 10px; margin: 5px 0; border-radius: 12px; box-shadow: 0 1px 8px rgba(0,0,0,0.1); max-width: 85%; word-wrap: break-word; animation: slideIn 0.3s ease-out; font-size: 13px; position: relative; cursor: pointer; }
        .msg.selecting { padding-left: 35px; }
        .msg.selected { background-color: rgba(77, 208, 183, 0.2) !important; border: 2px solid #4dd0b7; }
        .msg-checkbox { position: absolute; left: 8px; top: 50%; transform: translateY(-50%); display: none; }
        .selecting .msg-checkbox { display: block; }
        .msg-checkbox input { width: 16px; height: 16px; }
        .user { background: linear-gradient(135deg, #4dd0b7 0%, #2bc77a 100%); color: white; align-self: flex-end; margin-left: auto; position: relative; }
        .user::after { content: ''; position: absolute; right: -6px; bottom: 6px; width: 0; height: 0; border: 6px solid transparent; border-left: 6px solid #2bc77a; }
        .bot { background: white; color: #333; align-self: flex-start; margin-right: auto; border: 1px solid #e0e0e0; position: relative; }
        .bot::after { content: ''; position: absolute; left: -6px; bottom: 6px; width: 0; height: 0; border: 6px solid transparent; border-right: 6px solid white; }
        #messages { height: 300px; overflow-y: auto; padding: 10px; background: rgba(255,255,255,0.3); border-radius: 10px; margin-bottom: 10px; backdrop-filter: blur(10px); display: flex; flex-direction: column; }
        .search-panel { background: white; border-radius: 10px; margin-bottom: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); overflow: hidden; }
        .search-banner { background: linear-gradient(135deg, #4dd0b7 0%, #2bc77a 100%); color: white; padding: 10px 15px; cursor: pointer; display: flex; align-items: center; justify-content: space-between; transition: all 0.3s; }
        .search-banner:hover { background: linear-gradient(135deg, #2bc77a 0%, #4dd0b7 100%); }
        .search-banner-title { font-size: 13px; font-weight: 600; display: flex; align-items: center; }
        .search-banner-title::before { content: '🔍'; margin-right: 8px; font-size: 14px; }
        .search-toggle { font-size: 16px; transition: transform 0.3s; }
        .search-toggle.open { transform: rotate(180deg); }
        .search-content { padding: 0; max-height: 0; overflow: hidden; transition: all 0.3s ease-out; }
        .search-content.open { padding: 15px; max-height: 250px; }
        .search-row { display: flex; gap: 8px; margin-bottom: 10px; flex-wrap: wrap; }
        .search-field { flex: 1; min-width: 120px; }
        .search-field label { display: block; font-size: 10px; color: #555; margin-bottom: 3px; font-weight: 600; }
        .search-field input { width: 100%; padding: 6px 8px; border: 1px solid #e0e0e0; border-radius: 6px; transition: border-color 0.3s; font-size: 12px; }
        .search-field input:focus { border-color: #4dd0b7; outline: none; }
        .search-actions { text-align: center; margin-top: 8px; }
        .input-section { display: flex; gap: 5px; background: white; padding: 10px; border-radius: 20px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
        .input-section input { flex: 1; padding: 10px 12px; border: none; border-radius: 15px; outline: none; background: #f8f9fa; font-size: 13px; }
        .file-upload-btn { background: #6c5ce7; color: white; display: flex; align-items: center; justify-content: center; width: 40px; height: 40px; border-radius: 50%; border: none; cursor: pointer; transition: all 0.3s; }
        .file-upload-btn:hover { background: #5a4fcf; transform: scale(1.1); }
        .file-input { display: none; }
        .admin-protected { display: none; }
        .admin-user .admin-protected { display: block; }
        .input-section button { padding: 10px 15px; border: none; border-radius: 15px; cursor: pointer; font-weight: 600; transition: all 0.3s; font-size: 12px; }
        .send-btn { background: linear-gradient(135deg, #4dd0b7 0%, #2bc77a 100%); color: white; }
        .send-btn:hover { transform: translateY(-1px); box-shadow: 0 2px 8px rgba(45, 199, 122, 0.4); }
        .pdf-btn { background: linear-gradient(135deg, #ff6b6b 0%, #ee5a24 100%); color: white; }
        .pdf-btn:hover { transform: translateY(-1px); box-shadow: 0 2px 8px rgba(255, 107, 107, 0.4); }
        .clear-btn { background: linear-gradient(135deg, #6c5ce7 0%, #a29bfe 100%); color: white; padding: 6px 12px; border-radius: 15px; font-size: 10px; }
        .clear-btn:hover { transform: translateY(-1px); box-shadow: 0 2px 8px rgba(108, 92, 231, 0.4); }
        .typing { opacity: 0.7; }
        .sources { background: #f8f9fa; padding: 6px 8px; border-radius: 6px; margin-top: 6px; border-left: 3px solid #4dd0b7; font-size: 11px; }
        .typing { opacity: 0.8; font-style: italic; }
        .message-time { font-size: 9px; color: #888; margin-top: 3px; }
        @media (max-width: 768px) {
          .chat-container { padding: 5px; }
          .search-row { flex-direction: column; gap: 5px; }
          .search-field { min-width: 100%; }
          .input-section { gap: 3px; padding: 8px; }
          .msg { max-width: 95%; }
        }
        @keyframes slideIn { from { opacity: 0; transform: translateY(20px); } to { opacity: 1; transform: translateY(0); } }
      </style>
    </head>
    <body>
      <div class="chat-container">
        <div class="header">
          <div class="header-left">
            <img src="/static/images/kanoonpk-logo.jpg" alt="KanoonPK Logo" class="logo">
            <div class="header-text">
              <h1>AI Legal Research Assistant</h1>
              <p class="tagline">Your trusted partner for Pakistan law research and legal insights</p>
            </div>
          </div>
          <div class="header-right" style="display: flex; align-items: center; gap: 15px;">
            <div class="user-profile">
              <img src="https://ui-avatars.com/api/?name=Legal+User&background=4dd0b7&color=fff&size=35" alt="User Avatar" class="user-avatar" id="userAvatar">
              <div class="user-info">
                <p class="user-name" id="userName">Legal User</p>
                <p class="user-status">Subscriber</p>
              </div>
            </div>
            <a href="/admin" class="admin-link admin-protected">🔧 Admin Panel</a>
          </div>
        </div>
        
        <div class="chat-actions">
          <button class="chat-action-btn delete-chat-btn" onclick="clearChat()">🗑️ Clear Chat</button>
          <button class="chat-action-btn select-mode-btn" id="selectModeBtn" onclick="toggleSelectMode()">☑️ Select</button>
          
          <div class="export-dropdown">
            <button class="chat-action-btn" onclick="toggleExportMenu()" style="background: #17a2b8; color: white;">📤 Export</button>
            <div class="export-menu" id="exportMenu">
              <div class="export-option" onclick="exportChat('txt')">📄 Export as TXT</div>
              <div class="export-option" onclick="exportChat('pdf')">📁 Export as PDF</div>
              <div class="export-option" onclick="exportChat('json')">📊 Export as JSON</div>
              <div class="export-option" onclick="exportSelectedMessages()">☑️ Export Selected</div>
            </div>
          </div>
        </div>
        
        <div class="bulk-actions" id="bulkActions">
          <button class="bulk-action-btn delete-selected-btn" onclick="deleteSelected()">🗑️ Delete</button>
          <button class="bulk-action-btn forward-selected-btn" onclick="forwardSelected()">➤ Forward</button>
          <button class="bulk-action-btn cancel-selection-btn" onclick="cancelSelection()">✕ Cancel</button>
          <span id="selectedCount" style="font-size: 11px; color: #666; margin-left: 10px;">0 selected</span>
        </div>
        
        <div id="messages">
          <div class='msg bot' data-msg-id="welcome">
            <div class="msg-checkbox"><input type="checkbox" onchange="updateSelection()"></div>
            🙏 Welcome to KanoonPK! I'm your AI legal research assistant for Pakistan law. Ask me about legal cases, statutes, or upload documents for analysis.
            <div class="message-time">Online</div>
          </div>
        </div>
        
        <div class="search-panel">
          <div class="search-banner" onclick="toggleSearch()">
            <div class="search-banner-title">Advanced Search Filters</div>
            <div class="search-toggle" id="searchToggle">▼</div>
          </div>
          <div class="search-content" id="searchContent">
            <div class="search-row">
              <div class="search-field">
                <label for="citationFilter">📑 Citations:</label>
                <input type="text" id="citationFilter" placeholder="e.g., PLD 2020 SC">
              </div>
              <div class="search-field">
                <label for="yearFilter">📅 Year:</label>
                <input type="text" id="yearFilter" placeholder="e.g., 2020">
              </div>
              <div class="search-field">
                <label for="pageFilter">📄 Page:</label>
                <input type="text" id="pageFilter" placeholder="e.g., 123">
              </div>
              <div class="search-field">
                <label for="courtFilter">🏛️ Court:</label>
                <input type="text" id="courtFilter" placeholder="e.g., Supreme Court">
              </div>
            </div>
            <div class="search-actions">
              <button onclick="clearFilters()" class="clear-btn">🗑️ Clear Filters</button>
            </div>
          </div>
        </div>
        
        <div class="input-section">
          <input type="file" id="documentInput" class="file-input" accept=".pdf,.docx,.txt" onchange="handleFileUpload()">
          <button onclick="document.getElementById('documentInput').click()" class="file-upload-btn" title="Upload Document for Analysis">📁</button>
          <input id="userInput" placeholder="Ask about Pakistan law..." onkeydown="if(event.key==='Enter')sendMessage()">
          <button onclick="sendMessage()" class="send-btn">💬 Send</button>
          <button onclick="downloadPDF()" class="pdf-btn">📄 PDF</button>
        </div>
      </div>

      <script>
        let lastAnswer = "";
        let lastCitations = [];
        let isSelectMode = false;
        let messageIdCounter = 0;
        
        function getCurrentTime() {
          const now = new Date();
          return now.toLocaleTimeString([], {hour: '2-digit', minute:'2-digit'});
        }

        function addMessage(content, isUser = false, isTyping = false) {
          const msgBox = document.getElementById("messages");
          const messageDiv = document.createElement('div');
          const time = getCurrentTime();
          const msgId = 'msg_' + (++messageIdCounter);
          
          messageDiv.className = `msg ${isUser ? 'user' : 'bot'}${isTyping ? ' typing' : ''}${isSelectMode ? ' selecting' : ''}`;
          messageDiv.setAttribute('data-msg-id', msgId);
          
          const checkboxHtml = `<div class="msg-checkbox"><input type="checkbox" onchange="updateSelection()"></div>`;
          
          if (isTyping) {
            messageDiv.id = 'typing';
            messageDiv.innerHTML = checkboxHtml + content;
          } else {
            messageDiv.innerHTML = checkboxHtml + `${content}<div class="message-time">${time}</div>`;
          }
          
          msgBox.appendChild(messageDiv);
          
          // Smooth scroll to bottom
          setTimeout(() => {
            msgBox.scrollTop = msgBox.scrollHeight;
          }, 50);
          
          return messageDiv;
        }

        async function sendMessage() {
          const input = document.getElementById("userInput");
          const userText = input.value.trim();
          if (!userText) return;

          // Get filter values
          const filters = {
            citation: document.getElementById("citationFilter").value.trim(),
            year: document.getElementById("yearFilter").value.trim(),
            page: document.getElementById("pageFilter").value.trim(),
            court: document.getElementById("courtFilter").value.trim()
          };

          // Add user message with animation
          addMessage(userText, true);
          input.value = "";
          
          // Show typing indicator
          const typingMsg = addMessage('💭 KanoonPK is analyzing your query...', false, true);

          try {
            const res = await fetch("/chat", {
              method: "POST",
              headers: { "Content-Type": "application/json" },
              body: JSON.stringify({ message: userText, filters: filters })
            });
            
            if (!res.ok) {
              throw new Error('Network response was not ok');
            }
            
            const data = await res.json();
            
            // Remove typing indicator
            typingMsg.remove();

            lastAnswer = data.reply;
            lastCitations = data.sources;

            // Build bot response with proper formatting
            let content = data.reply.replace(/\\n/g,"<br>");
            
            if (data.sources.length) {
              content += `<div class='sources'><strong>📑 Legal Sources:</strong> ${data.sources.join(", ")}</div>`;
            }
            
            if (Object.values(filters).some(f => f)) {
              const activeFilters = Object.entries(filters).filter(([k, v]) => v).map(([k, v]) => `${k}: ${v}`).join(", ");
              content += `<div class='sources'><strong>🔍 Search Filters:</strong> ${activeFilters}</div>`;
            }
            
            addMessage(content);
            
          } catch (error) {
            console.error('Chat error:', error);
            if (typingMsg && typingMsg.parentNode) {
              typingMsg.remove();
            }
            addMessage('⚠️ Sorry, there was an error processing your request. Please try again.', false);
          }
        }

        function toggleSearch() {
          const searchContent = document.getElementById('searchContent');
          const searchToggle = document.getElementById('searchToggle');
          
          if (searchContent.classList.contains('open')) {
            searchContent.classList.remove('open');
            searchToggle.classList.remove('open');
          } else {
            searchContent.classList.add('open');
            searchToggle.classList.add('open');
          }
        }

        function clearFilters() {
          document.getElementById("citationFilter").value = "";
          document.getElementById("yearFilter").value = "";
          document.getElementById("pageFilter").value = "";
          document.getElementById("courtFilter").value = "";
        }
        
        function clearChat() {
          if (confirm('Are you sure you want to clear all messages? This action cannot be undone.')) {
            const msgBox = document.getElementById('messages');
            msgBox.innerHTML = `
              <div class='msg bot' data-msg-id="welcome">
                <div class="msg-checkbox"><input type="checkbox" onchange="updateSelection()"></div>
                🙏 Welcome to KanoonPK! I'm your AI legal research assistant for Pakistan law. Ask me about legal cases, statutes, or upload documents for analysis.
                <div class="message-time">Online</div>
              </div>
            `;
            messageIdCounter = 0;
            if (isSelectMode) {
              toggleSelectMode();
            }
          }
        }
        
        function toggleSelectMode() {
          const selectBtn = document.getElementById('selectModeBtn');
          const bulkActions = document.getElementById('bulkActions');
          const messages = document.querySelectorAll('.msg');
          
          isSelectMode = !isSelectMode;
          
          if (isSelectMode) {
            selectBtn.classList.add('active');
            selectBtn.innerHTML = '☑️ Selecting...';
            bulkActions.classList.add('active');
            messages.forEach(msg => {
              msg.classList.add('selecting');
              const checkbox = msg.querySelector('input[type="checkbox"]');
              if (checkbox) checkbox.checked = false;
            });
          } else {
            selectBtn.classList.remove('active');
            selectBtn.innerHTML = '☑️ Select';
            bulkActions.classList.remove('active');
            messages.forEach(msg => {
              msg.classList.remove('selecting', 'selected');
              const checkbox = msg.querySelector('input[type="checkbox"]');
              if (checkbox) checkbox.checked = false;
            });
          }
          
          updateSelection();
        }
        
        function updateSelection() {
          const selectedCount = document.querySelectorAll('.msg input[type="checkbox"]:checked').length;
          document.getElementById('selectedCount').textContent = selectedCount + ' selected';
          
          // Update visual selection
          document.querySelectorAll('.msg').forEach(msg => {
            const checkbox = msg.querySelector('input[type="checkbox"]');
            if (checkbox && checkbox.checked) {
              msg.classList.add('selected');
            } else {
              msg.classList.remove('selected');
            }
          });
        }
        
        function deleteSelected() {
          const selectedMessages = document.querySelectorAll('.msg input[type="checkbox"]:checked');
          if (selectedMessages.length === 0) {
            alert('Please select messages to delete.');
            return;
          }
          
          if (confirm(`Delete ${selectedMessages.length} selected message(s)?`)) {
            selectedMessages.forEach(checkbox => {
              const msgElement = checkbox.closest('.msg');
              msgElement.remove();
            });
            updateSelection();
          }
        }
        
        function forwardSelected() {
          const selectedMessages = document.querySelectorAll('.msg input[type="checkbox"]:checked');
          if (selectedMessages.length === 0) {
            alert('Please select messages to forward.');
            return;
          }
          
          let forwardText = 'Forwarded Messages:\n\n';
          selectedMessages.forEach(checkbox => {
            const msgElement = checkbox.closest('.msg');
            const msgContent = msgElement.textContent.replace(/☑️|Online|\d{1,2}:\d{2}/g, '').trim();
            const isUser = msgElement.classList.contains('user');
            forwardText += `${isUser ? 'You' : 'KanoonPK'}: ${msgContent}\n\n`;
          });
          
          // Copy to clipboard
          navigator.clipboard.writeText(forwardText).then(() => {
            alert(`${selectedMessages.length} message(s) copied to clipboard!`);
          }).catch(() => {
            // Fallback: show in alert
            prompt('Copy this text to forward:', forwardText);
          });
        }
        
        function cancelSelection() {
          if (isSelectMode) {
            toggleSelectMode();
          }
        }
        
        async function handleFileUpload() {
          const fileInput = document.getElementById('documentInput');
          const file = fileInput.files[0];
          
          if (!file) return;
          
          // Show upload progress
          const uploadMsg = addMessage(`📁 Uploading and analyzing: ${file.name}...`, false, true);
          
          const formData = new FormData();
          formData.append('file', file);
          
          try {
            const response = await fetch('/upload-and-analyze', {
              method: 'POST',
              body: formData
            });
            
            if (!response.ok) {
              throw new Error('Upload failed');
            }
            
            const result = await response.json();
            
            // Remove upload progress
            uploadMsg.remove();
            
            // Add analysis result
            addMessage(`📄 Document uploaded and indexed successfully!<br><strong>File:</strong> ${file.name}<br><strong>Analysis:</strong> ${result.summary || 'Document processed and ready for queries.'}`);
            
            // Clear file input
            fileInput.value = '';
            
          } catch (error) {
            console.error('Upload error:', error);
            uploadMsg.remove();
            addMessage('❌ Error uploading document. Please try again.', false);
            fileInput.value = '';
          }
        }
        
        // Check if user is admin (simple demo - in production use proper auth)
        function checkAdminStatus() {
          const userName = localStorage.getItem('userName') || 'Legal User';
          const isAdmin = localStorage.getItem('isAdmin') === 'true';
          
          document.getElementById('userName').textContent = userName;
          document.getElementById('userAvatar').src = `https://ui-avatars.com/api/?name=${encodeURIComponent(userName)}&background=4dd0b7&color=fff&size=35`;
          
          if (isAdmin) {
            document.body.classList.add('admin-user');
          }
        }
        
        function toggleExportMenu() {
          const menu = document.getElementById('exportMenu');
          menu.classList.toggle('show');
          
          // Close menu when clicking outside
          document.addEventListener('click', function(e) {
            if (!e.target.closest('.export-dropdown')) {
              menu.classList.remove('show');
            }
          });
        }
        
        function collectChatMessages(selectedOnly = false) {
          const messages = [];
          const msgElements = selectedOnly ? 
            document.querySelectorAll('.msg input[type="checkbox"]:checked') :
            document.querySelectorAll('.msg');
          
          (selectedOnly ? 
            Array.from(msgElements).map(cb => cb.closest('.msg')) :
            Array.from(msgElements)
          ).forEach(msg => {
            const isUser = msg.classList.contains('user');
            const isBot = msg.classList.contains('bot');
            const timeElement = msg.querySelector('.message-time');
            const time = timeElement ? timeElement.textContent : getCurrentTime();
            
            // Extract message content (excluding checkboxes and timestamps)
            let content = msg.textContent
              .replace(/\s*✅\s*/, '')
              .replace(time, '')
              .trim();
            
            // Remove checkbox artifacts
            content = content.replace(/^\s*/, '');
            
            if (content && !msg.classList.contains('typing')) {
              messages.push({
                type: isUser ? 'user' : 'bot',
                content: content,
                timestamp: time,
                id: msg.getAttribute('data-msg-id') || 'unknown'
              });
            }
          });
          
          return messages;
        }
        
        async function exportChat(format) {
          document.getElementById('exportMenu').classList.remove('show');
          
          const messages = collectChatMessages();
          if (messages.length === 0) {
            alert('No messages to export.');
            return;
          }
          
          try {
            const response = await fetch('/export-chat', {
              method: 'POST',
              headers: { 'Content-Type': 'application/json' },
              body: JSON.stringify({ messages: messages, format: format })
            });
            
            if (!response.ok) {
              throw new Error('Export failed');
            }
            
            const blob = await response.blob();
            const url = window.URL.createObjectURL(blob);
            const a = document.createElement('a');
            
            const timestamp = new Date().toISOString().slice(0, 19).replace(/:/g, '-');
            const filename = `KanoonPK_Chat_${timestamp}.${format}`;
            
            a.href = url;
            a.download = filename;
            document.body.appendChild(a);
            a.click();
            a.remove();
            window.URL.revokeObjectURL(url);
            
            addMessage(`✅ Chat exported as ${format.toUpperCase()} successfully!`, false);
            
          } catch (error) {
            console.error('Export error:', error);
            addMessage('❌ Error exporting chat. Please try again.', false);
          }
        }
        
        async function exportSelectedMessages() {
          document.getElementById('exportMenu').classList.remove('show');
          
          const selectedMessages = collectChatMessages(true);
          if (selectedMessages.length === 0) {
            alert('Please select messages to export.');
            return;
          }
          
          const format = prompt('Export format (txt, pdf, json):', 'txt');
          if (!format || !['txt', 'pdf', 'json'].includes(format.toLowerCase())) {
            return;
          }
          
          try {
            const response = await fetch('/export-chat', {
              method: 'POST',
              headers: { 'Content-Type': 'application/json' },
              body: JSON.stringify({ messages: selectedMessages, format: format.toLowerCase() })
            });
            
            if (!response.ok) {
              throw new Error('Export failed');
            }
            
            const blob = await response.blob();
            const url = window.URL.createObjectURL(blob);
            const a = document.createElement('a');
            
            const timestamp = new Date().toISOString().slice(0, 19).replace(/:/g, '-');
            const filename = `KanoonPK_Selected_${timestamp}.${format}`;
            
            a.href = url;
            a.download = filename;
            document.body.appendChild(a);
            a.click();
            a.remove();
            window.URL.revokeObjectURL(url);
            
            addMessage(`✅ Selected messages exported as ${format.toUpperCase()} successfully!`, false);
            
          } catch (error) {
            console.error('Export error:', error);
            addMessage('❌ Error exporting selected messages. Please try again.', false);
          }
        }
        
        async function downloadPDF() {
          if (!lastAnswer) {
            alert("⚠️ No answer to download yet.");
            return;
          }
          try {
            const res = await fetch("/export_pdf", {
              method: "POST",
              headers: { "Content-Type": "application/json" },
              body: JSON.stringify({ text: lastAnswer, citations: lastCitations })
            });
            
            if (!res.ok) {
              throw new Error('Failed to generate PDF');
            }
            
            const blob = await res.blob();
            const url = window.URL.createObjectURL(blob);
            const a = document.createElement("a");
            a.href = url;
            a.download = "KanoonPK_Answer.pdf";
            document.body.appendChild(a);
            a.click();
            a.remove();
            window.URL.revokeObjectURL(url);
          } catch (error) {
            console.error('PDF download error:', error);
            alert('⚠️ Error generating PDF. Please try again.');
          }
        }
        
        // Initialize on page load
        document.addEventListener('DOMContentLoaded', function() {
          checkAdminStatus();
          
          // Demo: Allow users to set their name and admin status
          document.getElementById('userName').addEventListener('click', function() {
            const newName = prompt('Enter your name:', this.textContent);
            if (newName && newName.trim()) {
              localStorage.setItem('userName', newName.trim());
              const makeAdmin = confirm('Are you an admin user?');
              localStorage.setItem('isAdmin', makeAdmin.toString());
              checkAdminStatus();
            }
          });
        });
      </script>
    </body>
    </html>
    """)

@app.route("/admin", methods=["GET", "POST"])
def admin():
    # Simple admin check - in production, implement proper authentication
    # For demo purposes, we'll assume access is controlled client-side
    if request.method == "POST":
        file = request.files["file"]
        if file and file.filename:
            filename = secure_filename(file.filename)
            path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            file.save(path)
            
            # Get additional metadata from form
            citation = request.form.get("citation", "")
            year = request.form.get("year", "")
            page = request.form.get("page", "")
            court = request.form.get("court", "")
            
            save_to_collection(path, filename, citation, year, page, court)
            return f"<div class='success'>✅ {filename} uploaded & indexed successfully!</div><a href='/admin' style='color: #4dd0b7; text-decoration: none; font-weight: 600;'>Upload Another Document</a>"
    
    # Get existing documents for display
    documents = get_uploaded_documents()
    
    # Build documents HTML
    documents_html = ""
    if documents:
        documents_html = "<div class='documents-grid'>"
        for doc in documents:
            doc_details = ""
            if doc['citation']:
                doc_details += f"<p><strong>Citation:</strong> {doc['citation']}</p>"
            if doc['year']:
                doc_details += f"<p><strong>Year:</strong> {doc['year']}</p>"
            if doc['page']:
                doc_details += f"<p><strong>Page:</strong> {doc['page']}</p>"
            if doc['court']:
                doc_details += f"<p><strong>Court:</strong> {doc['court']}</p>"
            doc_details += f"<p><strong>Chunks:</strong> {doc['chunks']}</p>"
            
            documents_html += f"""
            <div class="document-card">
                <div class="doc-header">
                    <h4>{doc['filename']}</h4>
                    <button onclick="deleteDocument('{doc['filename']}')" class="delete-btn" title="Delete Document">🗑️</button>
                </div>
                <div class="doc-details">
                    {doc_details}
                </div>
            </div>
            """
        documents_html += "</div>"
    else:
        documents_html = "<div class='no-documents'><p>No documents uploaded yet. Upload your first document below!</p></div>"
    
    return render_template_string("""
    <!DOCTYPE html>
    <html>
    <head>
        <title>KanoonPK Admin Panel</title>
        <style>
            * { margin: 0; padding: 0; box-sizing: border-box; }
            body { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%); min-height: 100vh; padding: 20px; }
            .container { max-width: 700px; margin: 0 auto; background: white; padding: 40px; border-radius: 20px; box-shadow: 0 10px 40px rgba(0,0,0,0.1); }
            .header { text-align: center; margin-bottom: 40px; }
            .logo { max-width: 150px; height: auto; margin-bottom: 15px; }
            .form-group { margin-bottom: 25px; }
            label { display: block; margin-bottom: 8px; font-weight: 600; color: #2bc77a; font-size: 14px; }
            input[type="text"], input[type="file"] { width: 100%; padding: 12px; border: 2px solid #e0e0e0; border-radius: 10px; font-size: 14px; transition: border-color 0.3s; }
            input[type="text"]:focus, input[type="file"]:focus { border-color: #4dd0b7; outline: none; }
            input[type="submit"] { background: linear-gradient(135deg, #4dd0b7 0%, #2bc77a 100%); color: white; padding: 15px 40px; border: none; border-radius: 25px; cursor: pointer; font-size: 16px; font-weight: 600; transition: all 0.3s; width: 100%; }
            input[type="submit"]:hover { transform: translateY(-2px); box-shadow: 0 8px 25px rgba(45, 199, 122, 0.4); }
            .file-types { font-size: 12px; color: #666; margin-top: 8px; padding: 8px; background: #f8f9fa; border-radius: 5px; }
            .title { color: #2bc77a; font-size: 24px; font-weight: 700; margin-bottom: 10px; }
            .subtitle { color: #666; margin-bottom: 30px; }
            .back-link { display: inline-block; margin-top: 30px; color: #4dd0b7; text-decoration: none; font-weight: 600; padding: 10px 25px; border: 2px solid #4dd0b7; border-radius: 20px; transition: all 0.3s; }
            .back-link:hover { background: #4dd0b7; color: white; }
            .success { color: #2bc77a; font-weight: bold; margin: 20px 0; padding: 15px; background: #d4edda; border-radius: 10px; border-left: 4px solid #2bc77a; }
            .documents-section { margin-bottom: 30px; }
            .documents-grid { display: grid; grid-template-columns: repeat(auto-fill, minmax(300px, 1fr)); gap: 20px; margin-bottom: 20px; }
            .document-card { background: #f8f9fa; border: 1px solid #e0e0e0; border-radius: 10px; padding: 15px; }
            .doc-header { display: flex; justify-content: space-between; align-items: center; margin-bottom: 10px; }
            .doc-header h4 { color: #2bc77a; margin: 0; font-size: 16px; }
            .delete-btn { background: #ff6b6b; color: white; border: none; border-radius: 50%; width: 30px; height: 30px; cursor: pointer; font-size: 14px; transition: all 0.3s; }
            .delete-btn:hover { background: #ee5a24; transform: scale(1.1); }
            .doc-details p { margin: 5px 0; font-size: 13px; color: #666; }
            .no-documents { text-align: center; color: #666; padding: 40px 20px; background: #f8f9fa; border-radius: 10px; }
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <img src="/static/images/kanoonpk-logo.jpg" alt="KanoonPK Logo" class="logo">
                <div class="title">Admin Panel</div>
                <div class="subtitle">Upload and manage legal documents</div>
            </div>
            
            <!-- Existing Documents Section -->
            <div class="documents-section">
                <h3 style="color: #2bc77a; margin-bottom: 20px; font-size: 20px;">📚 Uploaded Documents</h3>
""" + documents_html + """
            </div>
            
            <hr style="margin: 40px 0; border: none; height: 1px; background: #e0e0e0;">
            
            <!-- Upload Form Section -->
            <h3 style="color: #2bc77a; margin-bottom: 20px; font-size: 20px;">📤 Upload New Document</h3>
            <form method='POST' enctype='multipart/form-data'>
                <div class="form-group">
                    <label for="file">📁 Select Document:</label>
                    <input type='file' name='file' id='file' accept=".pdf,.docx,.txt,.jpg,.jpeg,.png" required>
                    <div class="file-types">Supported: PDF, DOCX, TXT, JPG, JPEG, PNG</div>
                </div>
                
                <div class="form-group">
                    <label for="citation">📑 Citation/Case Name:</label>
                    <input type='text' name='citation' id='citation' placeholder="e.g., PLD 2020 SC 123">
                </div>
                
                <div class="form-group">
                    <label for="year">📅 Year:</label>
                    <input type='text' name='year' id='year' placeholder="e.g., 2020">
                </div>
                
                <div class="form-group">
                    <label for="page">📄 Page Number:</label>
                    <input type='text' name='page' id='page' placeholder="e.g., 123">
                </div>
                
                <div class="form-group">
                    <label for="court">🏛️ Court:</label>
                    <input type='text' name='court' id='court' placeholder="e.g., Supreme Court of Pakistan">
                </div>
                
                <input type='submit' value='📤 Upload & Index Document'>
            </form>
            <a href='/' class="back-link">← Back to Chat</a>
        </div>
        
        <script>
            async function deleteDocument(filename) {
                if (!confirm('Are you sure you want to delete "' + filename + '"? This action cannot be undone.')) {
                    return;
                }
                
                try {
                    const response = await fetch('/admin/delete', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify({ filename: filename })
                    });
                    
                    const result = await response.json();
                    
                    if (result.success) {
                        alert('Document deleted successfully!');
                        location.reload();
                    } else {
                        alert('Error deleting document: ' + result.message);
                    }
                } catch (error) {
                    alert('Error deleting document: ' + error.message);
                }
            }
        </script>
    </body>
    </html>
    """)

@app.route("/chat", methods=["POST"])
def chat():
    user_input = request.json.get("message") if request.json else None
    filters = request.json.get("filters", {}) if request.json else {}
    
    if not user_input:
        return jsonify({"reply": "Please provide a message.", "sources": []})
    
    # Build where clause for filtering
    where_clause = {}
    if filters.get("citation"):
        where_clause["citation"] = {"$contains": filters["citation"]}
    if filters.get("year"):
        where_clause["year"] = {"$contains": filters["year"]}
    if filters.get("page"):
        where_clause["page"] = {"$contains": filters["page"]}
    if filters.get("court"):
        where_clause["court"] = {"$contains": filters["court"]}
    
    # Query with or without filters
    if where_clause:
        results = collection.query(
            query_texts=[user_input], 
            n_results=5,
            where=where_clause
        )
    else:
        results = collection.query(query_texts=[user_input], n_results=3)
    
    context = ""
    citations_used = []
    
    if results["documents"] and results["documents"][0]:
        for i, doc in enumerate(results["documents"][0]):
            metadata = results["metadatas"][0][i]
            citation = metadata.get("citation", metadata.get("source", "Unknown"))
            year = metadata.get("year", "")
            page = metadata.get("page", "")
            court = metadata.get("court", "")
            
            # Build citation display
            citation_display = citation
            if year or page or court:
                details = []
                if year: details.append(f"Year: {year}")
                if page: details.append(f"Page: {page}")
                if court: details.append(f"Court: {court}")
                citation_display += f" ({', '.join(details)})"
            
            citations_used.append(citation_display)
            context += f"\n\n[Citation: {citation_display}] {doc[:800]}..."
    
    # Use OpenAI with different prompts based on whether documents were found
    if context:
        # Documents found - use document-based response
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": LEGAL_SYSTEM_PROMPT},
                {"role": "user", "content": user_input},
                {"role": "assistant", "content": "Relevant documents:\n" + context}
            ]
        )
        reply = response.choices[0].message.content
    else:
        # No documents found - use general legal knowledge
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": GENERAL_LEGAL_PROMPT},
                {"role": "user", "content": user_input}
            ]
        )
        reply = response.choices[0].message.content + "\n\n📝 *Note: This answer is based on general knowledge of Pakistan law. For more specific information, please upload relevant legal documents.*"
        citations_used = ["General Legal Knowledge"]

    # Log history
    log_entry = {
        "timestamp": datetime.datetime.now().isoformat(),
        "question": user_input,
        "reply": reply,
        "citations": citations_used
    }
    os.makedirs("logs", exist_ok=True)
    history_file = "logs/history.json"
    if os.path.exists(history_file):
        with open(history_file, "r", encoding="utf-8") as f:
            history = json.load(f)
    else:
        history = []
    history.append(log_entry)
    with open(history_file, "w", encoding="utf-8") as f:
        json.dump(history, f, indent=2)

    return jsonify({"reply": reply, "sources": citations_used})

@app.route("/history")
def history():
    history_file = "logs/history.json"
    if not os.path.exists(history_file):
        return "<h2>No history yet.</h2>"

    with open(history_file, "r", encoding="utf-8") as f:
        history = json.load(f)

    html = "<h1>⚖️ KanoonPK — Search History</h1>"
    html += "<a href='/export_csv'>⬇️ Download as CSV</a><br><br>"
    html += "<table border='1' cellpadding='8' style='border-collapse:collapse;'>"
    html += "<tr><th>Time</th><th>Question</th><th>AI Reply (short)</th><th>Citations</th></tr>"

    for h in reversed(history[-50:]):
        short_reply = (h['reply'][:200] + "...") if len(h['reply']) > 200 else h['reply']
        html += f"<tr><td>{h['timestamp']}</td><td>{h['question']}</td><td>{short_reply}</td><td>{', '.join(h['citations'])}</td></tr>"

    html += "</table>"
    return html

@app.route("/export_csv")
def export_csv():
    history_file = "logs/history.json"
    if not os.path.exists(history_file):
        return "No history to export."
    with open(history_file, "r", encoding="utf-8") as f:
        history = json.load(f)

    filename = f"exports/history_{uuid.uuid4().hex}.csv"
    os.makedirs("exports", exist_ok=True)

    with open(filename, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(["Timestamp", "Question", "Reply", "Citations"])
        for h in history:
            writer.writerow([h["timestamp"], h["question"], h["reply"], "; ".join(h["citations"])])

    return send_file(filename, as_attachment=True)

@app.route("/upload-and-analyze", methods=["POST"])
def upload_and_analyze():
    """Handle file upload from chat interface"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No file provided'}), 400
    
    file = request.files['file']
    if not file or not file.filename:
        return jsonify({'success': False, 'message': 'No file selected'}), 400
    
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    
    try:
        file.save(file_path)
        
        # Process and add to collection
        save_to_collection(file_path, filename)
        
        # Generate summary
        if filename.lower().endswith('.pdf'):
            text_preview = extract_text_from_pdf(file_path)[:500] + "..."
        elif filename.lower().endswith('.docx'):
            text_preview = extract_text_from_docx(file_path)[:500] + "..."
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_preview = f.read()[:500] + "..."
        
        return jsonify({
            'success': True, 
            'message': 'File uploaded and indexed successfully',
            'filename': filename,
            'summary': f'Document contains legal text and has been indexed for search. Preview: {text_preview}'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Upload failed: {str(e)}'}), 500

@app.route("/admin/delete", methods=["POST"])
def admin_delete():
    data = request.json if request.json else {}
    filename = data.get('filename')
    
    if not filename:
        return jsonify({'success': False, 'message': 'No filename provided'})
    
    if delete_document(filename):
        return jsonify({'success': True, 'message': 'Document deleted successfully'})
    else:
        return jsonify({'success': False, 'message': 'Failed to delete document'})

@app.route("/export-chat", methods=["POST"])
def export_chat():
    """Export chat messages in various formats"""
    data = request.json if request.json else {}
    messages = data.get('messages', [])
    format_type = data.get('format', 'txt').lower()
    
    if not messages:
        return jsonify({'error': 'No messages provided'}), 400
    
    timestamp = datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    filename = f"exports/KanoonPK_Chat_{timestamp}.{format_type}"
    os.makedirs("exports", exist_ok=True)
    
    try:
        if format_type == 'txt':
            with open(filename, 'w', encoding='utf-8') as f:
                f.write("KanoonPK Chat Export\n")
                f.write("=" * 50 + "\n")
                f.write(f"Exported on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Total Messages: {len(messages)}\n\n")
                
                for msg in messages:
                    sender = "You" if msg['type'] == 'user' else "KanoonPK"
                    f.write(f"[{msg['timestamp']}] {sender}: {msg['content']}\n\n")
        
        elif format_type == 'json':
            export_data = {
                'export_info': {
                    'exported_at': datetime.datetime.now().isoformat(),
                    'total_messages': len(messages),
                    'format': 'json',
                    'source': 'KanoonPK AI Legal Research Assistant'
                },
                'messages': messages
            }
            
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
        
        elif format_type == 'pdf':
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            
            c = canvas.Canvas(filename, pagesize=A4)
            width, height = A4
            
            # Header
            c.setFont("Helvetica-Bold", 16)
            c.setFillColor(colors.HexColor("#2bc77a"))
            c.drawString(50, height - 50, "KanoonPK Chat Export")
            
            c.setFont("Helvetica", 10)
            c.setFillColor(colors.black)
            c.drawString(50, height - 70, f"Exported: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            c.drawString(50, height - 85, f"Messages: {len(messages)}")
            
            # Draw separator line
            c.line(50, height - 95, width - 50, height - 95)
            
            y = height - 120
            
            for msg in messages:
                sender = "You" if msg['type'] == 'user' else "KanoonPK"
                
                # Message header
                c.setFont("Helvetica-Bold", 10)
                c.setFillColor(colors.HexColor("#4dd0b7") if msg['type'] == 'user' else colors.HexColor("#6c757d"))
                c.drawString(50, y, f"[{msg['timestamp']}] {sender}:")
                y -= 15
                
                # Message content
                c.setFont("Helvetica", 9)
                c.setFillColor(colors.black)
                
                # Word wrap for long messages
                words = msg['content'].split(' ')
                line = ""
                for word in words:
                    test_line = line + word + " "
                    if c.stringWidth(test_line, "Helvetica", 9) < (width - 100):
                        line = test_line
                    else:
                        if line:
                            c.drawString(70, y, line.strip())
                            y -= 12
                        line = word + " "
                        
                        if y < 100:  # Start new page
                            c.showPage()
                            y = height - 50
                
                if line:
                    c.drawString(70, y, line.strip())
                    y -= 20
                
                if y < 100:  # Start new page
                    c.showPage()
                    y = height - 50
            
            c.save()
        
        else:
            return jsonify({'error': 'Unsupported format'}), 400
        
        return send_file(filename, as_attachment=True)
    
    except Exception as e:
        return jsonify({'error': f'Export failed: {str(e)}'}), 500

@app.route("/export_pdf", methods=["POST"])
def export_pdf():
    data = request.json
    text = data.get("text", "") if data else ""
    citations = data.get("citations", []) if data else []

    filename = f"exports/{uuid.uuid4().hex}.pdf"
    os.makedirs("exports", exist_ok=True)

    c = canvas.Canvas(filename, pagesize=A4)
    width, height = A4

    # Watermark
    c.setFont("Helvetica-Bold", 60)
    c.setFillGray(0.9, 0.3)
    c.saveState()
    c.translate(width/2, height/2)
    c.rotate(45)
    c.drawCentredString(0, 0, "KanoonPK")
    c.restoreState()

    y = height - 60

    # Citation Banner
    if citations:
        c.setFillColor(colors.HexColor("#007BFF"))
        c.rect(40, y-30, width-80, 30, fill=1)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y-20, "Citations: " + ", ".join(citations))
        y -= 50

    # Answer text
    c.setFillColor(colors.black)
    c.setFont("Helvetica", 12)
    for line in text.split("\n"):
        c.drawString(50, y, line)
        y -= 18
        if y < 80:
            c.showPage()
            y = height - 80
            c.setFont("Helvetica", 12)

    c.save()
    return send_file(filename, as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
